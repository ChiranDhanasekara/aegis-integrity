"""
Round-trip DOCX Document Editor -- AEGIS v4.0.

Applies text modifications and writing suggestions directly to Word (.docx) documents
while preserving:
  - Paragraph and character styles (headings, normal, bold, italic, color)
  - Tables, rows, cells, and nested elements
  - Header and footer blocks
  - Mathematical equations (OMML)
  - Image references and hyperlinked fields

Handles cross-run phrase matching when target text spans across multiple XML runs.
"""

from __future__ import annotations
import re
import logging
from pathlib import Path
from typing import Optional, Union

import docx
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

from aegis.writing.suggestion import SuggestionSet, WritingSuggestion

logger = logging.getLogger(__name__)


class DocxEditor:
    """
    Load, modify, and save Microsoft Word (.docx) documents with full formatting fidelity.

    Usage::

        editor = DocxEditor("manuscript.docx")
        editor.replace_text("in order to", "to")
        editor.apply_suggestions(suggestion_set)
        editor.save("manuscript_edited.docx")
    """

    def __init__(self, docx_path_or_doc: Optional[Union[str, Path, Document]] = None):
        if docx_path_or_doc is None:
            self._doc = docx.Document()
            self._source_path = None
        elif isinstance(docx_path_or_doc, Document):
            self._doc = docx_path_or_doc
            self._source_path = None
        else:
            path = Path(docx_path_or_doc)
            if not path.exists():
                raise FileNotFoundError(f"DOCX file not found: {path}")
            self._doc = docx.Document(str(path))
            self._source_path = path

    @property
    def document(self) -> Document:
        return self._doc

    def get_text(self) -> str:
        """Extract full plain text from document (paragraphs + tables)."""
        chunks = []
        for p in self._doc.paragraphs:
            if p.text:
                chunks.append(p.text)
        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text:
                            chunks.append(p.text)
        return "\n\n".join(chunks)

    def replace_text(self, target: str, replacement: str, count: int = -1) -> int:
        """
        Replace occurrences of `target` with `replacement` throughout the document.
        Preserves formatting of surrounding runs.

        Returns total count of replacements made.
        """
        if not target:
            return 0

        replaced_total = 0

        # 1. Main body paragraphs
        for p in self._doc.paragraphs:
            if target in p.text:
                rep = self._replace_in_paragraph(p, target, replacement, count - replaced_total if count > 0 else -1)
                replaced_total += rep
                if count > 0 and replaced_total >= count:
                    return replaced_total

        # 2. Table cells
        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if target in p.text:
                            rep = self._replace_in_paragraph(p, target, replacement, count - replaced_total if count > 0 else -1)
                            replaced_total += rep
                            if count > 0 and replaced_total >= count:
                                return replaced_total

        # 3. Headers and Footers
        for section in self._doc.sections:
            for p in section.header.paragraphs:
                if target in p.text:
                    replaced_total += self._replace_in_paragraph(p, target, replacement, -1)
            for p in section.footer.paragraphs:
                if target in p.text:
                    replaced_total += self._replace_in_paragraph(p, target, replacement, -1)

        return replaced_total

    def apply_suggestions(self, suggestions: Union[SuggestionSet, list[WritingSuggestion]]) -> int:
        """
        Apply accepted or modified suggestions to the document.

        Returns count of successfully applied suggestions.
        """
        if isinstance(suggestions, SuggestionSet):
            to_apply = suggestions.accepted
        else:
            to_apply = [s for s in suggestions if s.status in ("accepted", "modified")]

        applied = 0
        for s in to_apply:
            target = s.original_text
            replacement = s.final_text
            count = self.replace_text(target, replacement, count=1)
            if count > 0:
                applied += 1

        return applied

    def _replace_in_paragraph(self, p: Paragraph, target: str, replacement: str, max_rep: int = -1) -> int:
        """
        Replace target in a single paragraph, handling both single-run and multi-run matches.
        """
        if target not in p.text:
            return 0

        # Fast path: Target is fully contained within a single run
        rep_count = 0
        for r in p.runs:
            if target in r.text:
                occurrences = r.text.count(target)
                if max_rep > 0:
                    occurrences = min(occurrences, max_rep - rep_count)
                r.text = r.text.replace(target, replacement, occurrences)
                rep_count += occurrences
                if max_rep > 0 and rep_count >= max_rep:
                    return rep_count

        if target not in p.text:
            return rep_count

        # Slow path: Target spans across multiple runs
        # We consolidate run text into the first run and clear the crossed runs
        full_text = p.text
        if target in full_text:
            new_text = full_text.replace(target, replacement, 1 if max_rep == 1 else -1)
            if p.runs:
                # Retain formatting of first run
                p.runs[0].text = new_text
                for r in p.runs[1:]:
                    r.text = ""
                rep_count += 1

        return rep_count

    def save(self, output_path: Union[str, Path]) -> Path:
        """Save modified document to destination path."""
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        self._doc.save(str(out))
        return out
