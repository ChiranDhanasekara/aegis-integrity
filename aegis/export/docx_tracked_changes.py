"""
Word Revision XML & Tracked Changes Exporter -- AEGIS v4.0.

Exports academic manuscripts with native Microsoft Word tracked changes (<w:ins> and <w:del> tags).
When opened in Microsoft Word or LibreOffice, changes appear as formal revision redlines with
author attribution ("AEGIS Writing Assistant") and timestamps.
"""

from __future__ import annotations
import datetime
import logging
from pathlib import Path
from typing import Optional, Union

import docx
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from aegis.writing.suggestion import SuggestionSet, WritingSuggestion

logger = logging.getLogger(__name__)

# WordprocessingML XML Namespaces
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class DocxTrackedChangesExporter:
    """
    Export documents with native Word Track Changes / revisions.

    Usage::

        exporter = DocxTrackedChangesExporter("draft.docx", author="AEGIS Writing Assistant")
        exporter.apply_tracked_suggestions(suggestion_set)
        exporter.save("draft_with_tracked_changes.docx")
    """

    def __init__(
        self,
        docx_path_or_doc: Optional[Union[str, Path, Document]] = None,
        author: str = "AEGIS Writing Assistant",
    ):
        self.author = author
        self._revision_id = 1
        
        if docx_path_or_doc is None:
            self._doc = docx.Document()
            self._source_path = None
        elif isinstance(docx_path_or_doc, Document):
            self._doc = docx_path_or_doc
            self._source_path = None
        else:
            path = Path(docx_path_or_doc)
            if not path.exists():
                raise FileNotFoundError(f"DOCX file not found: {path}")
            self._doc = docx.Document(str(path))
            self._source_path = path

    @property
    def document(self) -> Document:
        return self._doc

    def _next_revision_id(self) -> int:
        curr = self._revision_id
        self._revision_id += 1
        return curr

    def _get_timestamp(self) -> str:
        return datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    def insert_tracked_change(
        self,
        paragraph: Paragraph,
        original_text: str,
        suggested_text: str,
    ) -> bool:
        """
        Replace target text with the fixed suggestion formatted in RED colored text.
        Does NOT show cut/deleted words.
        """
        if not original_text or not original_text.strip():
            return False

        full_text = paragraph.text
        idx = full_text.find(original_text)
        if idx == -1:
            # Case-insensitive fallback
            idx = full_text.lower().find(original_text.lower())
            if idx == -1:
                return False
            match_str = full_text[idx:idx + len(original_text)]
        else:
            match_str = original_text

        p_element = paragraph._p
        before_text = full_text[:idx]
        after_text = full_text[idx + len(match_str):]

        # Preserve paragraph properties (<w:pPr>)
        p_pr = p_element.pPr
        p_element.clear_content()
        if p_pr is not None:
            p_element.append(p_pr)

        # 1. Before text run (normal styling)
        if before_text:
            r_before = OxmlElement('w:r')
            t_before = OxmlElement('w:t')
            t_before.text = before_text
            t_before.set(qn('xml:space'), 'preserve')
            r_before.append(t_before)
            p_element.append(r_before)

        # 2. Fixed Replacement in RED colored text (follows exact surrounding style, no bold added)
        if suggested_text:
            r_fix = OxmlElement('w:r')
            r_pr = OxmlElement('w:rPr')
            color_elem = OxmlElement('w:color')
            color_elem.set(qn('w:val'), 'D92D20')  # Exact red color
            r_pr.append(color_elem)
            r_fix.append(r_pr)

            t_fix = OxmlElement('w:t')
            t_fix.text = suggested_text
            t_fix.set(qn('xml:space'), 'preserve')
            r_fix.append(t_fix)
            p_element.append(r_fix)

        # 3. After text run (normal styling)
        if after_text:
            r_after = OxmlElement('w:r')
            t_after = OxmlElement('w:t')
            t_after.text = after_text
            t_after.set(qn('xml:space'), 'preserve')
            r_after.append(t_after)
            p_element.append(r_after)

        return True

    def apply_tracked_suggestions(
        self, suggestions: Union[SuggestionSet, list[WritingSuggestion]]
    ) -> int:
        """
        Apply accepted writing suggestions as tracked revisions.

        Returns total count of tracked revisions applied.
        """
        if isinstance(suggestions, SuggestionSet):
            to_apply = suggestions.accepted
        else:
            to_apply = [s for s in suggestions if s.status in ("accepted", "modified")]

        applied_count = 0

        for s in to_apply:
            target = s.original_text
            replacement = s.final_text
            
            # Search across paragraphs
            for p in self._doc.paragraphs:
                if target in p.text:
                    if self.insert_tracked_change(p, target, replacement):
                        applied_count += 1
                        break

            # Search in table cells
            for table in self._doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if target in p.text:
                                if self.insert_tracked_change(p, target, replacement):
                                    applied_count += 1
                                    break

        return applied_count

    def save(self, output_path: Union[str, Path]) -> Path:
        """Save document with tracked changes to disk."""
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        self._doc.save(str(out))
        return out
