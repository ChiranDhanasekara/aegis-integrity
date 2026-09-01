"""
Unit tests for aegis.export.docx_editor and aegis.export.docx_tracked_changes.
"""

import tempfile
from pathlib import Path
import pytest
import docx

from aegis.export.docx_editor import DocxEditor
from aegis.export.docx_tracked_changes import DocxTrackedChangesExporter
from aegis.writing.suggestion import WritingSuggestion, SuggestionSet


class TestDocxEditor:

    def test_create_and_replace_text(self):
        editor = DocxEditor()
        editor.document.add_paragraph("This is an experiment in order to test the pipeline.")
        editor.document.add_paragraph("A second paragraph due to the fact that we need coverage.")
        
        replaced = editor.replace_text("in order to", "to")
        assert replaced == 1
        
        text = editor.get_text()
        assert "in order to" not in text
        assert "to test the pipeline" in text

    def test_apply_suggestions(self):
        editor = DocxEditor()
        editor.document.add_paragraph("The model is able to classify images rapidly.")
        
        sug = WritingSuggestion(
            category="wordiness",
            severity="info",
            original_text="is able to",
            suggested_text="can",
            explanation="Use can.",
            start_offset=10,
            end_offset=20,
            confidence=0.85,
        )
        sug.accept()
        
        applied = editor.apply_suggestions([sug])
        assert applied == 1
        assert "can classify images" in editor.get_text()

    def test_save_and_reload(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp_path = tmp.name

        try:
            editor = DocxEditor()
            editor.document.add_paragraph("Testing round-trip saving.")
            editor.save(tmp_path)
            
            reloaded = DocxEditor(tmp_path)
            assert "Testing round-trip saving." in reloaded.get_text()
        finally:
            if Path(tmp_path).exists():
                Path(tmp_path).unlink()


class TestDocxTrackedChangesExporter:

    def test_insert_tracked_change(self):
        exporter = DocxTrackedChangesExporter(author="AEGIS Reviewer")
        p = exporter.document.add_paragraph("We conducted the study in order to evaluate the accuracy.")
        
        success = exporter.insert_tracked_change(p, "in order to", "to")
        assert success is True
        
        # Verify XML contains red color highlight without cut words (no w:del)
        xml = p._p.xml
        assert "w:del" not in xml
        assert "in order to" not in xml
        assert "D92D20" in xml
        assert "to" in xml

    def test_apply_tracked_suggestions(self):
        exporter = DocxTrackedChangesExporter()
        exporter.document.add_paragraph("The dataset is comprised of 5000 images.")
        
        sug = WritingSuggestion(
            category="grammar",
            severity="error",
            original_text="comprised of",
            suggested_text="composed of",
            explanation="Usage error.",
            start_offset=15,
            end_offset=27,
            confidence=0.95,
        )
        sug.accept()
        
        applied = exporter.apply_tracked_suggestions([sug])
        assert applied == 1
        
        xml = exporter.document.paragraphs[0]._p.xml
        assert "w:del" not in xml
        assert "comprised of" not in xml
        assert "D92D20" in xml
        assert "composed of" in xml

    def test_save_tracked_changes(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp_path = tmp.name

        try:
            exporter = DocxTrackedChangesExporter()
            p = exporter.document.add_paragraph("Old phrasing to be replaced.")
            exporter.insert_tracked_change(p, "Old phrasing", "New revised phrasing")
            out_file = exporter.save(tmp_path)
            assert out_file.exists()
            assert out_file.stat().st_size > 0
        finally:
            if Path(tmp_path).exists():
                Path(tmp_path).unlink()
