"""
Regression test for the DOCX single-paragraph bug.

python-docx's Document.paragraphs already delineates each Word paragraph
correctly -- the bug was in how DocumentParser._parse_docx joined them:
a single "\n" leaves no blank line anywhere in full_text, but every
paragraph-level detector (AIContentDetector, NGramDetector) splits on
"\n\n+" to find paragraph boundaries. Real IEEE manuscripts tested with
the old code were reported as containing exactly one paragraph regardless
of actual length.
"""

from docx import Document as DocxDocument

from aegis.core.document import DocumentParser

PARAGRAPH_TEXTS = [
    f"This is paragraph number {i} of the manuscript, containing enough "
    f"distinct words to be a realistic academic sentence about network "
    f"security topic area number {i} for testing purposes."
    for i in range(1, 13)  # 12 paragraphs
]


def _make_docx(tmp_path, paragraphs, table_rows=None):
    doc = DocxDocument()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for i, row_vals in enumerate(table_rows):
            for j, val in enumerate(row_vals):
                table.cell(i, j).text = val
    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


class TestDocxParagraphPreservation:

    def test_full_text_has_blank_line_paragraph_separators(self, tmp_path):
        path = _make_docx(tmp_path, PARAGRAPH_TEXTS)
        parsed = DocumentParser().parse(str(path))
        # Splitting on the same delimiter downstream detectors use must
        # recover all 12 paragraphs, not collapse to 1.
        import re
        recovered = [p for p in re.split(r"\n\n+", parsed.full_text) if p.strip()]
        assert len(recovered) == len(PARAGRAPH_TEXTS)

    def test_ai_detector_paragraph_split_recovers_all_paragraphs(self, tmp_path):
        from aegis.detectors.ai_detector import AIContentDetector
        path = _make_docx(tmp_path, PARAGRAPH_TEXTS)
        parsed = DocumentParser().parse(str(path))
        det = AIContentDetector()
        # min_words default is 50; our paragraphs are short, so lower it to
        # match this test's fixture text length while still exercising the
        # real splitting logic end to end.
        paragraphs = det._split_paragraphs(parsed.full_text, min_words=5)
        assert len(paragraphs) >= 10, (
            f"expected ~12 paragraphs, got {len(paragraphs)} -- DOCX paragraphs "
            "are collapsing into fewer blocks than they should"
        )

    def test_ngram_detector_paragraph_split_recovers_all_paragraphs(self, tmp_path):
        from aegis.detectors.ngram import NGramDetector
        path = _make_docx(tmp_path, PARAGRAPH_TEXTS)
        parsed = DocumentParser().parse(str(path))
        det = NGramDetector()
        paragraphs = det._split_paragraphs(parsed.full_text, min_words=5)
        assert len(paragraphs) >= 10

    def test_table_content_is_included(self, tmp_path):
        rows = [["Metric", "Value"], ["Accuracy", "0.94"], ["Latency", "12ms"]]
        path = _make_docx(tmp_path, PARAGRAPH_TEXTS[:2], table_rows=rows)
        parsed = DocumentParser().parse(str(path))
        assert "Accuracy" in parsed.full_text
        assert "0.94" in parsed.full_text

    def test_no_regression_on_single_paragraph_doc(self, tmp_path):
        path = _make_docx(tmp_path, ["Just one short paragraph here."])
        parsed = DocumentParser().parse(str(path))
        assert "Just one short paragraph here." in parsed.full_text
