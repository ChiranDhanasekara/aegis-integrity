"""Tests for the mathematical formula checker."""

from aegis.detectors.math_formula import MathFormulaChecker


class TestTextHeuristicExtraction:

    def test_finds_numbered_equation_lines(self):
        text = "Intro\nE = mc^2 (1)\nF = ma (2)\n"
        result = MathFormulaChecker().analyze("paper.pdf", "pdf", text)
        assert result.equations_found == 2
        assert result.equation_numbers == ["1", "2"]

    def test_no_equations_found_in_plain_prose(self):
        text = "This paper discusses several ideas without any formulas at all."
        result = MathFormulaChecker().analyze("paper.pdf", "pdf", text)
        assert result.equations_found == 0
        assert result.numbering_issues == []


class TestNumberingChecks:

    def test_duplicate_numbers_flagged(self):
        text = "a = b (1)\nc = d (2)\ne = f (2)\n"
        result = MathFormulaChecker().analyze("paper.pdf", "pdf", text)
        messages = [i.message for i in result.numbering_issues]
        assert any("Duplicate" in m for m in messages)

    def test_gap_in_numbering_flagged(self):
        text = "a = b (1)\nc = d (3)\n"
        result = MathFormulaChecker().analyze("paper.pdf", "pdf", text)
        messages = [i.message for i in result.numbering_issues]
        assert any("gap" in m.lower() for m in messages)

    def test_clean_sequential_numbering_has_no_issues(self):
        text = "a = b (1)\nc = d (2)\ne = f (3)\n"
        result = MathFormulaChecker().analyze("paper.pdf", "pdf", text)
        assert result.numbering_issues == []


class TestReferenceChecks:

    def test_dangling_reference_flagged(self):
        text = "a = b (1)\nAs shown in (5), the result holds.\n"
        result = MathFormulaChecker().analyze("paper.pdf", "pdf", text)
        messages = [i.message for i in result.reference_issues]
        assert any("5" in m and "do not match" in m for m in messages)

    def test_valid_reference_not_flagged_as_dangling(self):
        text = "a = b (1)\nAs shown in (1), the result holds.\n"
        result = MathFormulaChecker().analyze("paper.pdf", "pdf", text)
        dangling = [i for i in result.reference_issues if "do not match" in i.message]
        assert dangling == []

    def test_chained_reference_resolves_both_numbers(self):
        text = "a = b (1)\nc = d (2)\nAs in (1) and (2) we conclude.\n"
        result = MathFormulaChecker().analyze("paper.pdf", "pdf", text)
        orphan_msgs = [i.message for i in result.reference_issues if "never referenced" in i.message]
        assert orphan_msgs == []

    def test_bare_number_without_cue_is_not_treated_as_a_reference(self):
        # Enumerated list items like "(1) first step (2) second step" must
        # not be misread as equation references -- this was the main
        # false-positive risk in the design (see module docstring).
        text = "a = b (1)\nThe steps are: (1) prepare, (2) execute, (3) verify.\n"
        result = MathFormulaChecker().analyze("paper.pdf", "pdf", text)
        dangling = [i for i in result.reference_issues if "do not match" in i.message]
        assert dangling == []

    def test_orphan_equation_flagged_but_low_severity(self):
        text = "a = b (1)\nc = d (2)\nAs shown in (1), the result holds.\n"
        result = MathFormulaChecker().analyze("paper.pdf", "pdf", text)
        orphan = [i for i in result.reference_issues if "never referenced" in i.message]
        assert len(orphan) == 1
        assert orphan[0].severity == "LOW"

    def test_inconsistent_phrasing_flagged(self):
        text = ("a = b (1)\nc = d (2)\n"
                "As shown in Eq. (1), and as in equation (2), we conclude.\n")
        result = MathFormulaChecker().analyze("paper.pdf", "pdf", text)
        messages = [i.message for i in result.reference_issues]
        assert any("inconsistent phrasing" in m for m in messages)


class TestLatexExtraction(object):

    def test_extracts_equation_environments(self, tmp_path):
        tex = tmp_path / "paper.tex"
        tex.write_text(
            r"""
            \begin{equation}
            E = mc^2
            \end{equation}
            \begin{equation}
            F = ma
            \end{equation}
            As in (1) and (2) we conclude.
            """,
            encoding="utf-8",
        )
        result = MathFormulaChecker().analyze(str(tex), "latex", "")
        assert result.equations_found == 2
        assert result.extraction_method == "latex_source"
        assert result.limitations  # numbers were inferred, not read

    def test_unbalanced_braces_flagged(self, tmp_path):
        tex = tmp_path / "paper.tex"
        tex.write_text(
            r"""
            \begin{equation}
            E = mc^{2
            \end{equation}
            """,
            encoding="utf-8",
        )
        result = MathFormulaChecker().analyze(str(tex), "latex", "")
        assert any(i.category == "malformed_latex" for i in result.notation_issues)


class TestDocxExtraction:

    def test_extracts_omml_equations(self, tmp_path):
        import docx
        from lxml import etree

        doc = docx.Document()
        doc.add_paragraph("Intro text.")
        p = doc.add_paragraph()
        omath_xml = (
            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
            "<m:r><m:t>E=mc</m:t></m:r></m:oMath>"
        )
        p._p.append(etree.fromstring(omath_xml))
        doc.add_paragraph("As shown in (1), energy and mass are related.")

        path = tmp_path / "paper.docx"
        doc.save(str(path))

        from aegis.core.document import DocumentParser
        parsed = DocumentParser().parse(str(path))
        result = MathFormulaChecker().analyze(str(path), "docx", parsed.full_text)

        assert result.extraction_method == "docx_omml"
        assert result.equations_found == 1

    def test_no_omml_present_notes_limitation(self, tmp_path):
        import docx
        doc = docx.Document()
        doc.add_paragraph("A paper with no equations at all.")
        path = tmp_path / "paper.docx"
        doc.save(str(path))

        from aegis.core.document import DocumentParser
        parsed = DocumentParser().parse(str(path))
        result = MathFormulaChecker().analyze(str(path), "docx", parsed.full_text)

        assert result.equations_found == 0
        assert any("No native Word equation objects" in lim for lim in result.limitations)


class TestNotationConventions:

    def test_exponential_notation_flagged(self):
        text = "The measured value was 5E03 ohms."
        result = MathFormulaChecker().analyze("paper.pdf", "pdf", text)
        assert any("exponential notation" in i.message for i in result.notation_issues)

    def test_percentage_range_missing_sign_flagged(self):
        text = "The students made up 20-30% of the population."
        result = MathFormulaChecker().analyze("paper.pdf", "pdf", text)
        assert any("percentage range" in i.message for i in result.notation_issues)

    def test_decimal_without_leading_zero_flagged(self):
        text = "The threshold was set to .25 for all trials."
        result = MathFormulaChecker().analyze("paper.pdf", "pdf", text)
        assert any("leading zero" in i.message for i in result.notation_issues)

    def test_clean_notation_has_no_issues(self):
        text = "The threshold was set to 0.25 across all 40-50 mm samples."
        result = MathFormulaChecker().analyze("paper.pdf", "pdf", text)
        assert result.notation_issues == []
